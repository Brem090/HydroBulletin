"""Універсальне формування трьох Word-бюлетенів із SQLite."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Mapping, Sequence

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from .archive import ProductResult, query_observations, register_product
from .quality import MISSING, VALID, worst_quality_status
from .regions import RegionConfig


PARAMETERS = (
    "WATER_LEVEL",
    "DAILY_CHANGE",
    "PRECIPITATION",
    "WATER_TEMPERATURE",
    "DISCHARGE",
)

QUALITY_DISPLAY = {
    "INCONSISTENT_CHANGE": "INCONSISTENT\nCHANGE",
    "OUT_OF_RANGE": "OUT OF\nRANGE",
}

PRECIPITATION_FONT_SIZE_PT = 10

UKRAINIAN_MONTHS_GENITIVE = (
    "",
    "січня",
    "лютого",
    "березня",
    "квітня",
    "травня",
    "червня",
    "липня",
    "серпня",
    "вересня",
    "жовтня",
    "листопада",
    "грудня",
)


@dataclass(frozen=True)
class BulletinRow:
    station_index: str
    station_name: str
    level: float | None
    change: float | None
    precipitation: float | None
    water_temperature: float | None
    discharge: float | None
    quality_status: str
    quality_message: str
    precipitation_source: str
    observation_ids: tuple[int, ...]


@dataclass(frozen=True)
class BulletinResult:
    region_key: str
    output_path: Path
    rows: tuple[BulletinRow, ...]
    product: ProductResult


def _format_number(value: float | None, digits: int = 1) -> str:
    if value is None:
        return ""
    if float(value).is_integer():
        return str(int(value))
    return f"{value:.{digits}f}".rstrip("0").rstrip(".").replace(".", ",")


def _format_change(value: float | None) -> str:
    if value is None:
        return ""
    text = _format_number(value, 1)
    return f"+{text}" if value > 0 else text


def _pick_measurement(
    candidates: Sequence[dict[str, object]],
    parameter_code: str,
) -> dict[str, object] | None:
    matching = [
        row for row in candidates if row["parameter_code"] == parameter_code
    ]
    if not matching:
        return None
    if parameter_code == "WATER_LEVEL":
        return min(
            matching,
            key=lambda row: abs(
                datetime.fromisoformat(str(row["observed_at"])).hour - 8
            ),
        )
    return max(matching, key=lambda row: str(row["observed_at"]))


def build_bulletin_rows(
    db_path: Path,
    region: RegionConfig,
    bulletin_date: str,
    precipitation_mapping: Mapping[str, str],
) -> tuple[BulletinRow, ...]:
    """Формує регіональну вибірку з урахуванням опадів SYNOP."""

    hydro_indexes = [station.index for station in region.stations]
    meteo_indexes = [
        precipitation_mapping[index]
        for index in hydro_indexes
        if index in precipitation_mapping
    ]
    records = query_observations(
        db_path,
        start_date=bulletin_date,
        end_date=bulletin_date,
        station_indexes=tuple(dict.fromkeys(hydro_indexes + meteo_indexes)),
        parameter_codes=PARAMETERS,
    )
    grouped: dict[str, list[dict[str, object]]] = {}
    for record in records:
        grouped.setdefault(str(record["station_index"]), []).append(record)

    result: list[BulletinRow] = []
    for station in region.stations:
        station_records = grouped.get(station.index, [])
        selected: dict[str, dict[str, object] | None] = {
            parameter: _pick_measurement(station_records, parameter)
            for parameter in PARAMETERS
        }

        hydro_precip = selected["PRECIPITATION"]
        meteo_index = precipitation_mapping.get(station.index)
        meteo_precip: dict[str, object] | None = None
        if meteo_index:
            meteo_precip = _pick_measurement(
                grouped.get(meteo_index, []),
                "PRECIPITATION",
            )

        precipitation_record = meteo_precip or hydro_precip
        selected["PRECIPITATION"] = precipitation_record

        present_records = [row for row in selected.values() if row is not None]
        statuses = [str(row["quality_status"]) for row in present_records]
        missing_fields: list[str] = []
        if selected["WATER_LEVEL"] is None:
            missing_fields.append("рівень")
        if selected["DAILY_CHANGE"] is None:
            missing_fields.append("добова зміна")

        quality_status = worst_quality_status(statuses)
        quality_messages = [
            str(row["quality_message"])
            for row in present_records
            if str(row["quality_message"]).strip()
        ]
        if missing_fields:
            quality_status = MISSING
            quality_messages.append(
                f"Відсутні обов'язкові поля: {', '.join(missing_fields)}."
            )
        elif quality_status in {"OK", "NOT_CHECKED"}:
            quality_status = VALID

        if meteo_precip is not None and meteo_index:
            precipitation_source = f"метеостанція {meteo_index}"
        elif hydro_precip is not None:
            precipitation_source = "гідропост"
        elif meteo_index:
            precipitation_source = f"метеостанція {meteo_index}: немає даних"
        else:
            precipitation_source = "немає даних"

        def value(parameter: str) -> float | None:
            row = selected[parameter]
            return None if row is None or row["value"] is None else float(row["value"])

        observation_ids = tuple(
            sorted(
                {
                    int(row["observation_id"])
                    for row in present_records
                    if row.get("observation_id") is not None
                }
            )
        )
        result.append(
            BulletinRow(
                station_index=station.index,
                station_name=station.name,
                level=value("WATER_LEVEL"),
                change=value("DAILY_CHANGE"),
                precipitation=value("PRECIPITATION"),
                water_temperature=value("WATER_TEMPERATURE"),
                discharge=value("DISCHARGE"),
                quality_status=quality_status,
                quality_message=" ".join(dict.fromkeys(quality_messages)),
                precipitation_source=precipitation_source,
                observation_ids=observation_ids,
            )
        )
    return tuple(result)


def _iter_tables(container):
    """Обходить також вкладені таблиці офіційних шаблонів."""

    seen_tables: set[object] = set()

    def walk(owner):
        for table in owner.tables:
            table_key = table._tbl
            if table_key in seen_tables:
                continue
            seen_tables.add(table_key)
            yield table

            seen_cells: set[object] = set()
            for row in table.rows:
                for cell in row.cells:
                    cell_key = cell._tc
                    if cell_key in seen_cells:
                        continue
                    seen_cells.add(cell_key)
                    yield from walk(cell)

    yield from walk(container)


def _iter_paragraphs(doc: Document):
    """Повертає абзаци документа, включно з усіма рівнями таблиць."""

    seen_paragraphs: set[object] = set()
    for paragraph in doc.paragraphs:
        paragraph_key = paragraph._p
        if paragraph_key not in seen_paragraphs:
            seen_paragraphs.add(paragraph_key)
            yield paragraph

    for table in _iter_tables(doc):
        seen_cells: set[object] = set()
        for row in table.rows:
            for cell in row.cells:
                cell_key = cell._tc
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                for paragraph in cell.paragraphs:
                    paragraph_key = paragraph._p
                    if paragraph_key in seen_paragraphs:
                        continue
                    seen_paragraphs.add(paragraph_key)
                    yield paragraph


def _set_paragraph_text(paragraph, text: str) -> None:
    """Замінює текст, зберігаючи формат видимого run шаблону."""

    runs = paragraph.runs
    if runs:
        target_run = next(
            (run for run in runs if run.text.strip()),
            runs[0],
        )
        target_run.text = text
        for run in runs:
            if run is not target_run:
                run.text = ""
    else:
        paragraph.add_run(text)


def _replace_markers(doc: Document, replacements: Mapping[str, str]) -> None:
    for paragraph in _iter_paragraphs(doc):
        full_text = "".join(run.text for run in paragraph.runs)
        replaced = full_text
        for marker, value in replacements.items():
            replaced = replaced.replace(marker, value)
        if replaced != full_text:
            _set_paragraph_text(paragraph, replaced)


def _format_ukrainian_date(bulletin_date: str) -> str:
    try:
        value = datetime.strptime(bulletin_date, "%d.%m.%Y")
    except ValueError as error:
        raise ValueError(
            f"Дата бюлетеня має формат ДД.ММ.РРРР: {bulletin_date}"
        ) from error
    month = UKRAINIAN_MONTHS_GENITIVE[value.month]
    return f"{value.day:02d} {month} {value.year} року"


def _replace_official_labels(
    doc: Document,
    *,
    bulletin_date: str,
    hydrologist: str,
) -> None:
    """Оновлює дату й підпис без додавання службових маркерів у DOCX."""

    date_found = False
    hydrologist_found = False
    formatted_date = _format_ukrainian_date(bulletin_date)
    for paragraph in _iter_paragraphs(doc):
        text = paragraph.text.strip()
        if text.startswith("станом на "):
            _set_paragraph_text(paragraph, f"станом на {formatted_date}")
            date_found = True
        elif text.startswith("Черговий гідролог:"):
            suffix = f" {hydrologist}" if hydrologist else ""
            _set_paragraph_text(paragraph, f"Черговий гідролог:{suffix}")
            hydrologist_found = True

    missing: list[str] = []
    if not date_found:
        missing.append("рядок дати «станом на ...»")
    if not hydrologist_found:
        missing.append("рядок «Черговий гідролог:»")
    if missing:
        raise RuntimeError(
            "Офіційний шаблон не містить обов'язкових елементів: "
            + ", ".join(missing)
            + "."
        )


def _set_cell_text(
    cell,
    text: str,
    *,
    color: RGBColor | None = None,
    font_size_pt: float | None = None,
) -> None:
    paragraph = cell.paragraphs[0]
    runs = paragraph.runs
    if runs:
        target_run = next(
            (run for run in runs if run.text.strip()),
            runs[0],
        )
        target_run.text = text
        for run in runs:
            if run is not target_run:
                run.text = ""
    else:
        target_run = paragraph.add_run(text)
    if color is not None:
        target_run.font.color.rgb = color
    if font_size_pt is not None:
        target_run.font.size = Pt(font_size_pt)


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _station_rows(table) -> dict[str, object]:
    result: dict[str, object] = {}
    for row in table.rows[1:]:
        if not row.cells:
            continue
        index = row.cells[0].text.strip()
        if index:
            result[index] = row
    return result


def _find_bulletin_table(doc: Document):
    """Визначає офіційну або розширену структуру шаблону."""

    extended_table = None
    for table in _iter_tables(doc):
        if not table.rows or not table.columns:
            continue
        first_header = table.rows[0].cells[0].text.strip()
        if len(table.columns) >= 10 and "Річка-пункт" in first_header:
            return "official", table
        if len(table.columns) >= 9 and first_header == "Індекс":
            extended_table = table

    if extended_table is not None:
        return "extended", extended_table
    raise RuntimeError(
        "У шаблоні не знайдено ані офіційної 10-стовпцевої таблиці, "
        "ані розширеної таблиці HydroBulletin."
    )


def _fill_official_table(table, rows: Sequence[BulletinRow]) -> None:
    """Заповнює лише змінні поля, не торкаючись порогів і багаторічних даних."""

    expected_rows = len(rows) + 2
    if len(table.rows) != expected_rows:
        raise RuntimeError(
            "Кількість рядків офіційного шаблону не відповідає довіднику: "
            f"очікується {expected_rows}, отримано {len(table.rows)}."
        )

    for position, item in enumerate(rows, start=2):
        target = table.rows[position]
        values = {
            1: _format_number(item.level, 1),
            2: _format_change(item.change),
            3: _format_number(item.precipitation, 1),
            9: _format_number(item.water_temperature, 1),
        }
        for column, text in values.items():
            _set_cell_text(
                target.cells[column],
                text,
                font_size_pt=(
                    PRECIPITATION_FONT_SIZE_PT if column == 3 else None
                ),
            )


def _fill_extended_table(table, rows: Sequence[BulletinRow]) -> None:
    """Зберігає сумісність із розширеним шаблоном HydroBulletin."""

    by_index = _station_rows(table)
    warning_color = RGBColor(192, 0, 0)
    for position, item in enumerate(rows, start=1):
        row = by_index.get(item.station_index)
        if row is None:
            row_index = position
            if row_index >= len(table.rows):
                row = table.add_row()
            else:
                row = table.rows[row_index]

        values = (
            item.station_index,
            item.station_name,
            _format_number(item.level, 1),
            _format_change(item.change),
            _format_number(item.precipitation, 1),
            _format_number(item.water_temperature, 1),
            _format_number(item.discharge, 3),
            QUALITY_DISPLAY.get(item.quality_status, item.quality_status),
            item.precipitation_source,
        )
        for column, text in enumerate(values):
            _set_cell_text(row.cells[column], text)

        if item.quality_status != VALID:
            _set_cell_text(
                row.cells[7],
                QUALITY_DISPLAY.get(item.quality_status, item.quality_status),
                color=warning_color,
            )
            _shade_cell(row.cells[7], "FCE8E6")


def generate_bulletin(
    db_path: Path,
    region: RegionConfig,
    *,
    bulletin_date: str,
    hydrologist: str,
    template_path: Path,
    output_path: Path,
    precipitation_mapping: Mapping[str, str],
) -> BulletinResult:
    """Заповнює Word-шаблон і зберігає походження даних у SQLite."""

    template_path = Path(template_path)
    if not template_path.exists():
        raise FileNotFoundError(f"Шаблон бюлетеня не знайдено: {template_path}")

    rows = build_bulletin_rows(
        db_path,
        region,
        bulletin_date,
        precipitation_mapping,
    )
    doc = Document(str(template_path))
    try:
        template_layout, table = _find_bulletin_table(doc)
    except RuntimeError as error:
        raise RuntimeError(f"Помилка шаблону {template_path.name}: {error}") from error

    if template_layout == "official":
        _replace_official_labels(
            doc,
            bulletin_date=bulletin_date,
            hydrologist=hydrologist,
        )
        _fill_official_table(table, rows)
    else:
        _replace_markers(
            doc,
            {
                "{{REGION}}": region.title,
                "{{DATE}}": bulletin_date,
                "{{HYDROLOGIST}}": hydrologist,
            },
        )
        _fill_extended_table(table, rows)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    observation_ids = {
        observation_id
        for row in rows
        for observation_id in row.observation_ids
    }
    quality_counts: dict[str, int] = {}
    for row in rows:
        quality_counts[row.quality_status] = (
            quality_counts.get(row.quality_status, 0) + 1
        )
    product = register_product(
        db_path,
        product_type="WORD_BULLETIN",
        region_key=region.key,
        bulletin_date=bulletin_date,
        output_path=str(output_path),
        observation_ids=observation_ids,
        metadata={
            "template": str(template_path),
            "hydrologist": hydrologist,
            "rows": len(rows),
            "precipitation_mapping_applied": True,
            "template_layout": template_layout,
            "quality_counts": quality_counts,
        },
    )
    return BulletinResult(region.key, output_path, rows, product)


def generate_bulletins(
    db_path: Path,
    regions: Iterable[RegionConfig],
    *,
    bulletin_date: str,
    hydrologist: str,
    templates_dir: Path,
    output_dir: Path,
    precipitation_mapping: Mapping[str, str],
) -> tuple[BulletinResult, ...]:
    """Створює вибрані бюлетені одним універсальним кодом."""

    results: list[BulletinResult] = []
    for region in regions:
        results.append(
            generate_bulletin(
                db_path,
                region,
                bulletin_date=bulletin_date,
                hydrologist=hydrologist,
                template_path=region.template_path(templates_dir),
                output_path=Path(output_dir) / region.output_name(bulletin_date),
                precipitation_mapping=precipitation_mapping,
            )
        )
    return tuple(results)
