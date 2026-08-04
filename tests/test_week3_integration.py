"""Інтеграційні перевірки пакетного імпорту, SQLite та DOCX."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from hydrobulletin.archive import (
    archive_summary,
    read_product_provenance,
)
from hydrobulletin.batch import discover_batch_files, run_batch_import
from hydrobulletin.regions import REGIONS
from hydrobulletin.stations import ALL_STATIONS, METEO_STATIONS, STATIONS_BY_INDEX
from hydrobulletin.workflow import WorkflowRequest, execute_workflow


PROJECT_DIR = Path(__file__).resolve().parents[1]
DEMO_DIR = PROJECT_DIR / "demo_data" / "week3"


def iter_tables(container):
    for table in container.tables:
        yield table
        seen_cells: set[object] = set()
        for row in table.rows:
            for cell in row.cells:
                if cell._tc in seen_cells:
                    continue
                seen_cells.add(cell._tc)
                yield from iter_tables(cell)


def official_data_table(document: Document):
    return next(
        table
        for table in iter_tables(document)
        if len(table.columns) == 10
        and "Річка-пункт" in table.rows[0].cells[0].text
    )


class BatchDiscoveryTests(unittest.TestCase):
    def test_discovers_files_in_chronological_order(self) -> None:
        files = discover_batch_files(DEMO_DIR)

        self.assertEqual(len(files), 3)
        self.assertEqual(files[0].bulletin_date, "11.07.2026")
        self.assertEqual(files[-1].message_type, "SYNOP")


class OfficialTemplateTests(unittest.TestCase):
    def test_templates_match_region_station_counts(self) -> None:
        templates_dir = PROJECT_DIR / "templates" / "bulletins"
        for region in REGIONS:
            with self.subTest(region=region.key):
                document = Document(str(region.template_path(templates_dir)))
                table = official_data_table(document)
                self.assertEqual(len(table.rows), len(region.stations) + 2)


class WeekThreeWorkflowTests(unittest.TestCase):
    def test_full_demo_creates_three_traceable_bulletins(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            db_path = root / "archive.sqlite"
            raw_root = root / "raw"
            output_dir = root / "output"

            batch = run_batch_import(
                DEMO_DIR,
                raw_root=raw_root,
                db_path=db_path,
                all_stations=ALL_STATIONS,
                hydro_stations_by_index=STATIONS_BY_INDEX,
                meteo_stations_by_index={
                    station.index: station for station in METEO_STATIONS
                },
            )
            self.assertEqual(batch.processed_files, 3)
            self.assertEqual(batch.errors, ())

            request = WorkflowRequest(
                bulletin_date="12.07.2026",
                source_mode="database",
                message_types=("ZRUR52",),
                db_path=db_path,
                raw_root=raw_root,
                templates_dir=PROJECT_DIR / "templates" / "bulletins",
                output_dir=output_dir,
                mapping_path=PROJECT_DIR
                / "config"
                / "precipitation_mapping.json",
                region_keys=("lviv", "if", "left_dnister"),
                hydrologist="Тестовий гідролог",
                include_meteo=False,
            )
            result = execute_workflow(request)

            self.assertEqual(len(result.bulletins), 3)
            self.assertEqual(
                result.quality_summary.counts["INCONSISTENT_CHANGE"],
                1,
            )
            lviv = next(
                item for item in result.bulletins if item.region_key == "lviv"
            )
            observation_with_meteo = next(
                row for row in lviv.rows if row.station_index == "79726"
            )
            self.assertEqual(observation_with_meteo.precipitation, 3.0)
            self.assertEqual(
                observation_with_meteo.precipitation_source,
                "метеостанція 33288",
            )

            left = next(
                item
                for item in result.bulletins
                if item.region_key == "left_dnister"
            )
            flagged = next(
                row for row in left.rows if row.station_index == "81244"
            )
            self.assertEqual(flagged.quality_status, "INCONSISTENT_CHANGE")

            for bulletin in result.bulletins:
                self.assertTrue(bulletin.output_path.exists())
                document = Document(str(bulletin.output_path))
                visible_text = "\n".join(
                    [paragraph.text for paragraph in document.paragraphs]
                    + [
                        cell.text
                        for table in document.tables
                        for row in table.rows
                        for cell in row.cells
                    ]
                )
                self.assertNotIn("{{", visible_text)
                self.assertIn("Тестовий гідролог", visible_text)
                self.assertIn("станом на 12 липня 2026 року", visible_text)

                data_table = official_data_table(document)
                self.assertEqual(
                    len(data_table.rows),
                    len(bulletin.rows) + 2,
                )

                if bulletin.region_key == "lviv":
                    mapped_row = data_table.rows[26]
                    self.assertIn("Кам.-Бузька", mapped_row.cells[0].text)
                    self.assertEqual(mapped_row.cells[3].text.strip(), "3")
                    precipitation_run = next(
                        run
                        for paragraph in mapped_row.cells[3].paragraphs
                        for run in paragraph.runs
                        if run.text.strip()
                    )
                    self.assertEqual(precipitation_run.font.size.pt, 10.0)
                    self.assertEqual(
                        [
                            mapped_row.cells[index].text.strip()
                            for index in range(4, 9)
                        ],
                        ["200", "330", "348", "117", "21"],
                    )
                    self.assertEqual(mapped_row.cells[9].text.strip(), "18")
                    temperature_run = next(
                        run
                        for paragraph in mapped_row.cells[9].paragraphs
                        for run in paragraph.runs
                        if run.text.strip()
                    )
                    self.assertEqual(
                        str(temperature_run.font.color.rgb),
                        "000000",
                    )

            provenance = read_product_provenance(
                db_path,
                lviv.product.product_id,
            )
            provenance_indexes = {
                str(row["station_index"]) for row in provenance
            }
            self.assertIn("79726", provenance_indexes)
            self.assertIn("33288", provenance_indexes)
            self.assertTrue(all(row["raw_path"] for row in provenance))

            # Повторне формування оновлює той самий продукт і його зв'язки.
            second = execute_workflow(request)
            self.assertEqual(len(second.bulletins), 3)
            self.assertEqual(archive_summary(db_path)["products"], 3)


if __name__ == "__main__":
    unittest.main()
