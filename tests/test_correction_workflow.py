"""Наскрізний сценарій правки, екстремумів, льоду та Word-бюлетеня."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import Table

from hydrobulletin.archive import (
    create_correction,
    initialize_archive,
    query_observations,
    read_product_provenance,
    upsert_reference_extreme,
)
from hydrobulletin.bulletins import generate_bulletin
from hydrobulletin.extremes import seed_extremes_from_templates
from hydrobulletin.pipeline import run_import_pipeline
from hydrobulletin.quality import run_initial_quality_control
from hydrobulletin.regions import REGIONS, REGIONS_BY_KEY
from hydrobulletin.sources import LocalFileSource
from hydrobulletin.stations import ALL_STATIONS, STATIONS_BY_INDEX


PROJECT_DIR = Path(__file__).resolve().parents[1]
TEMPLATES = PROJECT_DIR / "templates" / "bulletins"


def official_table(document: DocumentObject) -> Table:
    queue = list(document.tables)
    seen: set[object] = set()
    while queue:
        table = queue.pop(0)
        if table._tbl in seen:
            continue
        seen.add(table._tbl)
        if len(table.columns) >= 10 and "Річка-пункт" in table.rows[0].cells[0].text:
            return table
        for row in table.rows:
            for cell in row.cells:
                queue.extend(cell.tables)
    raise AssertionError("Офіційну таблицю не знайдено")


class CorrectionWorkflowTests(unittest.TestCase):
    def test_active_correction_and_reference_data_reach_word_and_provenance(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            db_path = root / "archive.sqlite"
            initialize_archive(db_path, ALL_STATIONS)
            seed_extremes_from_templates(db_path, REGIONS, TEMPLATES)

            raw_path = root / "12.07.2026_ZRUR52.txt"
            raw_path.write_text(
                "81015 12081 10105 20005 30103 40120 51605 70250 81234 00081 =\n",
                encoding="utf-8",
            )
            run_import_pipeline(
                LocalFileSource(raw_path),
                bulletin_date="12.07.2026",
                message_type="ZRUR52",
                raw_root=root / "raw",
                db_path=db_path,
                stations=ALL_STATIONS,
                stations_by_index=STATIONS_BY_INDEX,
                source_type="local",
                source_name=str(raw_path),
            )
            run_initial_quality_control(db_path, "12.07.2026")
            level = query_observations(
                db_path,
                start_date="12.07.2026",
                end_date="12.07.2026",
                station_indexes=("81015",),
                parameter_codes=("WATER_LEVEL",),
            )[0]
            observation_id = level["observation_id"]
            self.assertIsInstance(observation_id, int)
            assert isinstance(observation_id, int)
            correction = create_correction(
                db_path,
                observation_id,
                111,
                reason="Перевірено за журналом",
                hydrologist="Євген КОЗИРЄВ",
            )
            upsert_reference_extreme(
                db_path,
                station_index="81015",
                maximum_level=999,
                average_level=150,
                minimum_level=30,
                updated_by="Євген КОЗИРЄВ",
            )

            output_path = root / "bulletin.docx"
            result = generate_bulletin(
                db_path,
                REGIONS_BY_KEY["lviv"],
                bulletin_date="12.07.2026",
                hydrologist="Євген КОЗИРЄВ",
                template_path=TEMPLATES / "bulletin_lviv_template.docx",
                output_path=output_path,
                precipitation_mapping={},
            )

            table = official_table(Document(str(output_path)))
            self.assertEqual(table.rows[2].cells[1].text.strip(), "111")
            self.assertEqual(table.rows[2].cells[6].text.strip(), "999")
            self.assertIn("Льодохід 50%", table.rows[2].cells[9].text)
            self.assertIn("25 см", table.rows[2].cells[9].text)
            self.assertIn("льодові явища", table.rows[0].cells[9].text)

            provenance = read_product_provenance(
                db_path,
                result.product.product_id,
            )
            corrected = [
                row
                for row in provenance
                if row["correction_id"] == correction.correction_id
            ]
            self.assertEqual(len(corrected), 1)
            self.assertEqual(corrected[0]["original_value"], 105.0)
            self.assertEqual(corrected[0]["value"], 111.0)


if __name__ == "__main__":
    unittest.main()
